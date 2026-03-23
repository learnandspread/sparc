"""Semantic chunker for SPARC knowledge base documents.

Chunks .docx files by section headers rather than fixed-size windows,
preserving semantic coherence. Each chunk gets metadata for filtering.
"""

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document


@dataclass
class Chunk:
    text: str
    metadata: dict = field(default_factory=dict)


# Map source filenames to doc identifiers and content types
DOC_CONFIG = {
    "SPARC™ Master Knowledge System.docx": {
        "source_doc": "master_knowledge_system",
        "content_type": "doctrine",
    },
    "THE DISCOVERY & VISIBILITY ENCYCLOPEDIA.docx": {
        "source_doc": "discovery_encyclopedia",
        "content_type": "definition",
    },
    "SPARC™ RESPONSE TRAINING EXAMPLES.docx": {
        "source_doc": "response_training",
        "content_type": "example",
    },
}

# Keywords that map content to signal categories
SIGNAL_KEYWORDS = {
    "visibility": [
        "seo", "geo", "aeo", "search", "ranking", "impressions", "discovery",
        "google business profile", "gbp", "local search", "visibility", "indexation",
        "crawl", "serp", "keyword", "voice search", "schema", "structured data",
        "map pack", "local listing", "citation",
    ],
    "authority": [
        "authority", "e-e-a-t", "eeat", "backlink", "content depth", "pr",
        "digital pr", "entity", "topical", "expertise", "thought leadership",
        "content strategy", "pillar page", "cluster", "semantic depth",
    ],
    "trust": [
        "trust", "review", "reputation", "testimonial", "social proof",
        "credibility", "branding", "brand identity", "consistency", "case study",
    ],
    "engagement": [
        "engagement", "social media", "video", "email", "ux", "ui",
        "dwell time", "bounce", "interaction", "click-through", "ctr",
        "session", "scroll", "heatmap",
    ],
    "conversion": [
        "conversion", "cro", "landing page", "funnel", "cta", "form",
        "lead", "acquisition", "cac", "roas", "paid", "ads", "ppc",
        "retargeting", "campaign",
    ],
    "momentum": [
        "momentum", "growth", "velocity", "trend", "compound", "retention",
        "clv", "lvr", "lifecycle", "automation", "crm", "email marketing",
        "nurture", "loyalty",
    ],
}


def detect_signal_categories(text: str) -> list[str]:
    """Detect which signal categories a chunk of text relates to."""
    text_lower = text.lower()
    categories = []
    for category, keywords in SIGNAL_KEYWORDS.items():
        if any(kw in text_lower for kw in keywords):
            categories.append(category)
    return categories if categories else ["general"]


def extract_text_from_docx(filepath: str) -> str:
    """Extract all text from a .docx file."""
    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure with markdown-style markers
            if para.style and para.style.name and "Heading" in para.style.name:
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    level_num = int(level)
                    text = "#" * level_num + " " + text
                except ValueError:
                    text = "## " + text
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                table_text.append(row_text)
        if table_text:
            paragraphs.append("\n".join(table_text))

    return "\n\n".join(paragraphs)


def chunk_by_sections(text: str, min_chunk_size: int = 200, max_chunk_size: int = 2000) -> list[str]:
    """Split text into chunks by section headers.

    Uses markdown-style headers (# ## ###) as split points.
    Merges very small sections together.
    Splits very large sections into smaller pieces.
    """
    # Split on headers (keep the header with its content)
    header_pattern = r"(?=^#{1,4}\s)"
    raw_sections = re.split(header_pattern, text, flags=re.MULTILINE)
    raw_sections = [s.strip() for s in raw_sections if s.strip()]

    chunks = []
    buffer = ""

    for section in raw_sections:
        if len(buffer) + len(section) < min_chunk_size:
            buffer = (buffer + "\n\n" + section).strip()
        else:
            if buffer:
                chunks.append(buffer)
            buffer = section

    if buffer:
        chunks.append(buffer)

    # Split oversized chunks
    final_chunks = []
    for chunk in chunks:
        if len(chunk) > max_chunk_size:
            # Split by double newline, then regroup
            paragraphs = chunk.split("\n\n")
            sub_buffer = ""
            for para in paragraphs:
                if len(sub_buffer) + len(para) > max_chunk_size and sub_buffer:
                    final_chunks.append(sub_buffer.strip())
                    sub_buffer = para
                else:
                    sub_buffer = (sub_buffer + "\n\n" + para).strip()
            if sub_buffer:
                final_chunks.append(sub_buffer.strip())
        else:
            final_chunks.append(chunk)

    return final_chunks


def extract_section_name(chunk_text: str) -> str:
    """Extract the first header from a chunk as its section name."""
    for line in chunk_text.split("\n"):
        line = line.strip()
        if line.startswith("#"):
            return line.lstrip("#").strip()
    # Fallback: first 60 chars
    return chunk_text[:60].replace("\n", " ").strip()


def chunk_document(filepath: str) -> list[Chunk]:
    """Chunk a single document into metadata-tagged chunks."""
    path = Path(filepath)
    filename = path.name

    config = DOC_CONFIG.get(filename, {
        "source_doc": path.stem.lower().replace(" ", "_"),
        "content_type": "doctrine",
    })

    text = extract_text_from_docx(filepath)
    raw_chunks = chunk_by_sections(text)

    chunks = []
    for raw in raw_chunks:
        signal_cats = detect_signal_categories(raw)
        section = extract_section_name(raw)

        chunks.append(Chunk(
            text=raw,
            metadata={
                "source_doc": config["source_doc"],
                "section": section,
                "signal_category": signal_cats,
                "content_type": config["content_type"],
            },
        ))

    return chunks


def chunk_all_documents(filepaths: list[str]) -> list[Chunk]:
    """Chunk all knowledge base documents."""
    all_chunks = []
    for fp in filepaths:
        path = Path(fp)
        if not path.exists():
            print(f"Warning: {fp} not found, skipping")
            continue
        doc_chunks = chunk_document(fp)
        print(f"  {path.name}: {len(doc_chunks)} chunks")
        all_chunks.extend(doc_chunks)
    print(f"Total: {len(all_chunks)} chunks")
    return all_chunks
